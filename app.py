import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os
import datetime
import copy

# ---------------------------
# Config inicial
# ---------------------------
st.set_page_config(page_title="Gerador de Relatórios", layout="centered")
st.title("📑 Gerador de Relatórios Automático")

# ========================
# INICIALIZAÇÃO DO SESSION STATE (DEVE VIR PRIMEIRO)
# ========================
if "contrapartidas" not in st.session_state:
    st.session_state.contrapartidas = []
if "edit_index" not in st.session_state:
    st.session_state.edit_index = None

# ========================
# FORMULÁRIO PRINCIPAL
# ========================
nome_projeto = st.text_input("Nome do Projeto")
nome_sindicato = st.text_input("Nome do Sindicato")
cidade = st.text_input("Cidade")
data = st.date_input("Data", datetime.date.today())

# ========================
# CONTRAPARTIDAS
# ========================
st.subheader("➕ Adicionar Contrapartida")

col1, col2 = st.columns([3, 1])
with col1:
    nova_desc = st.text_input("Descrição da Contrapartida")
with col2:
    nova_status = st.selectbox("Comprovada?", ["Sim", "Não"])

if st.button("Adicionar Contrapartida"):
    if nova_desc:
        st.session_state.contrapartidas.append({
            "descricao": nova_desc,
            "status": nova_status
        })
        st.success(f"Contrapartida adicionada: {nova_desc} ({nova_status})")
        st.rerun()   # 🔄 reinicia app para limpar campo

# Listagem das contrapartidas (AGORA COM VERIFICAÇÃO CORRETA)
if len(st.session_state.contrapartidas) > 0:  # ✅ Verificação segura
    st.write("📋 Contrapartidas adicionadas:")
    for i, c in enumerate(st.session_state.contrapartidas):
        colA, colB, colC = st.columns([4, 1, 1])
        with colA:
            st.write(f"**{i+1}.** {c['descricao']} — *{c['status']}*")
        with colB:
            if st.button("✏️ Editar", key=f"edit_{i}"):
                st.session_state.edit_index = i
                st.rerun()
        with colC:
            if st.button("🗑 Remover", key=f"del_{i}"):
                st.session_state.contrapartidas.pop(i)
                st.rerun()

    # Modo edição
    if st.session_state.edit_index is not None:
        idx = st.session_state.edit_index
        if 0 <= idx < len(st.session_state.contrapartidas):
            st.info(f"✏️ Editando contrapartida {idx+1}")
            edit_desc = st.text_input(
                "Nova descrição",
                value=st.session_state.contrapartidas[idx]["descricao"]
            )
            edit_status = st.selectbox(
                "Comprovada?",
                ["Sim", "Não"],
                index=0 if st.session_state.contrapartidas[idx]["status"] == "Sim" else 1
            )
            colSave, colCancel = st.columns(2)
            with colSave:
                if st.button("💾 Salvar Alterações"):
                    st.session_state.contrapartidas[idx] = {
                        "descricao": edit_desc,
                        "status": edit_status
                    }
                    st.session_state.edit_index = None
                    st.rerun()
            with colCancel:
                if st.button("❌ Cancelar Edição"):
                    st.session_state.edit_index = None
                    st.rerun()
        else:
            st.session_state.edit_index = None

# ========================
# UPLOAD DE IMAGENS
# ========================
st.subheader("📷 Upload de Fotos")
imagens = st.file_uploader(
    "Selecione imagens para inserir no relatório",
    type=["jpg", "jpeg", "png", "bmp", "gif"],
    accept_multiple_files=True
)

compact_option = st.checkbox("Compactar imagens antes de inserir", value=True)

# ========================
# GERAR RELATÓRIO
# ========================
if st.button("Gerar Relatório"):
    if not (nome_projeto and nome_sindicato and cidade and data):
        st.error("⚠️ Por favor, preencha todos os campos obrigatórios!")
    else:
        try:
            doc = Document("Modelo.docx")

            # Normalizar textos
            nome_projeto_fmt = nome_projeto.upper()
            nome_sindicato_fmt = nome_sindicato.upper()
            cidade_fmt = cidade.capitalize()
            data_fmt = data.strftime("%d/%m/%Y")

            substituicoes = {
                "(TEXTO_NOMEPROJETO)": nome_projeto_fmt,
                "(TEXTO_NOMESINDICATO)": nome_sindicato_fmt,
                "(CIDADE)": cidade_fmt,
                "(DATA)": data_fmt
            }

            # Substituir em parágrafos e tabelas
            def substituir(doc, antigo, novo):
                for p in doc.paragraphs:
                    if antigo in p.text:
                        for run in p.runs:
                            run.text = run.text.replace(antigo, novo)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if antigo in p.text:
                                    for run in p.runs:
                                        run.text = run.text.replace(antigo, novo)

            for k, v in substituicoes.items():
                substituir(doc, k, v)

            # ========================
            # CONTRAPARTIDAS NA TABELA (VERSÃO CORRIGIDA)
            # ========================
            if len(st.session_state.contrapartidas) > 0:
                for table in doc.tables:
                    for row_idx, row in enumerate(table.rows):
                        row_text = ' '.join([cell.text for cell in row.cells])
                        
                        # Verificar se esta linha contém marcadores de contrapartida
                        for idx, c in enumerate(st.session_state.contrapartidas, start=1):
                            marcador = f"(contrapartida{idx:02d})"
                            
                            if marcador in row_text:
                                # Encontrar e substituir a descrição
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            if marcador in run.text:
                                                run.text = run.text.replace(marcador, c["descricao"])
                                
                                # Substituir checkboxes - abordagem mais específica
                                for cell_idx, cell in enumerate(row.cells):
                                    cell_text = cell.text
                                    
                                    # Se for célula da coluna SIM (normalmente 3ª coluna)
                                    if "(XSIM)" in cell_text:
                                        for p in cell.paragraphs:
                                            for run in p.runs:
                                                if c["status"] == "Sim":
                                                    run.text = run.text.replace("(XSIM)", "SIM").replace("(XNAO)", "")
                                                else:
                                                    run.text = run.text.replace("(XSIM)", "").replace("(XNAO)", "NÃO")
                                    
                                    # Se for célula da coluna NÃO (normalmente 4ª coluna)  
                                    if "(XNAO)" in cell_text:
                                        for p in cell.paragraphs:
                                            for run in p.runs:
                                                if c["status"] == "Não":
                                                    run.text = run.text.replace("(XNAO)", "NÃO").replace("(XSIM)", "")
                                                else:
                                                    run.text = run.text.replace("(XNAO)", "").replace("(XSIM)", "SIM")

                # Se houver extras, duplicar última linha
                max_default = 3
                if len(st.session_state.contrapartidas) > max_default:
                    for table in doc.tables:
                        if any("(contrapartida01)" in c.text for r in table.rows for c in r.cells):
                            template_row = table.rows[max_default]._tr
                            for extra_idx, c in enumerate(st.session_state.contrapartidas[max_default:], start=max_default+1):
                                new_tr = copy.deepcopy(template_row)
                                table._tbl.append(new_tr)
                                new_row = table.rows[-1]
                                
                                # Atualizar descrição
                                for cell in new_row.cells:
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            if "(contrapartida03)" in run.text:
                                                run.text = run.text.replace("(contrapartida03)", c["descricao"])
                                
                                # Atualizar checkboxes SIM/NÃO
                                for cell in new_row.cells:
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            if "(XSIM)" in run.text or "(XNAO)" in run.text:
                                                if c["status"] == "Sim":
                                                    run.text = run.text.replace("(XSIM)", "SIM").replace("(XNAO)", "")
                                                else:
                                                    run.text = run.text.replace("(XNAO)", "NÃO").replace("(XSIM)", "")
                            break

            # ========================
            # INSERIR IMAGENS
            # ========================
            for i, paragraph in enumerate(doc.paragraphs):
                if "(FOTOS_ORGANIZADAS)" in paragraph.text:
                    p_element = paragraph._element
                    p_element.getparent().remove(p_element)

                    if imagens:
                        for idx, img in enumerate(imagens, 1):
                            if compact_option:
                                with Image.open(img) as im:
                                    im = im.convert("RGB")
                                    max_width = 1600
                                    if im.width > max_width:
                                        ratio = max_width / im.width
                                        new_size = (max_width, int(im.height * ratio))
                                        im = im.resize(new_size, Image.LANCZOS)
                                    temp_path = f"temp_{idx}.jpg"
                                    im.save(temp_path, "JPEG", optimize=True, quality=70)
                            else:
                                temp_path = f"temp_{idx}_{img.name}"
                                with open(temp_path, "wb") as f:
                                    f.write(img.getbuffer())

                            if idx > 1:
                                doc.add_page_break()

                            p_img = doc.add_paragraph()
                            p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run_img = p_img.add_run()
                            run_img.add_picture(temp_path, width=Inches(6))

                            legenda = os.path.splitext(os.path.basename(img.name))[0]
                            p_legenda = doc.add_paragraph()
                            p_legenda.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run_legenda = p_legenda.add_run(f"Foto {idx}: {legenda}")
                            run_legenda.font.name = "Calibri"
                            run_legenda.font.size = Pt(10)
                            run_legenda.bold = True

                            try:
                                os.remove(temp_path)
                            except Exception:
                                pass
                    break

            # ========================
            # SALVAR DOCUMENTO
            # ========================
            output_path = "Relatorio_Gerado.docx"
            doc.save(output_path)

            with open(output_path, "rb") as f:
                st.success("✅ Relatório gerado com sucesso!")
                st.download_button(
                    "⬇️ Baixar Relatório",
                    f,
                    file_name="Relatorio_Gerado.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"❌ Erro ao gerar relatório: {e}")
